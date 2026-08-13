from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "LexRead_新版法学智能阅读器_PRD与实施清单.md"
OUTPUT = ROOT / "docs" / "LexRead_法律检索课程作业_PRD.docx"

INK = "17365D"
BLUE = "2E74B5"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
MUTED = "666666"


def set_font(run, name="Arial Unicode MS", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for i, width in enumerate(widths):
        grid.gridCol_lst[i].set(qn("w:w"), str(width))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color=MUTED)


def setup_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.86)
    sec.right_margin = Inches(0.86)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, INK, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.text = "LexRead 法学智能阅读与研究检索工作台｜产品需求文档"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_font(run, size=8.5, color=MUTED)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)


def inline_runs(paragraph, text, size=10.5):
    """Render minimal Markdown emphasis without turning prose into plain text."""
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=INK)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Menlo", size=size - 0.5, color=INK)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size)


def add_body_para(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.0 if indent else 0.22)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    inline_runs(p, text)


def add_bullet(doc, text, numbered=None):
    p = doc.add_paragraph(style="List Bullet" if numbered is None else "List Number")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    inline_runs(p, text)


def table_widths(count):
    presets = {2: [2700, 6660], 3: [1800, 3780, 3780], 4: [1320, 2700, 2700, 2640], 5: [1200, 2040, 2040, 2040, 2040]}
    return presets.get(count, [9360 // count] * count)


def add_table(doc, rows):
    if not rows:
        return
    count = len(rows[0])
    table = doc.add_table(rows=0, cols=count)
    table.style = "Table Grid"
    for r_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            inline_runs(p, value, size=8.7)
            if r_index == 0:
                shade(cell, LIGHT)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
            elif r_index % 2 == 0:
                shade(cell, "FAFBFC")
    set_table_geometry(table, table_widths(count))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("《法律检索》课程作业")
    set_font(run, name="Arial Unicode MS", size=16, color=MUTED)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LexRead 法学智能阅读与\n研究检索工作台")
    set_font(run, name="Arial Unicode MS", size=25, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("产品需求文档（PRD）")
    set_font(run, name="Arial Unicode MS", size=16, color=BLUE)
    p.paragraph_format.space_after = Pt(34)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["课程名称", "任课教师", "学生姓名 / 学号", "提交日期"]
    values = ["法律检索（请填写）", "请填写", "请填写", "2026 年 __ 月 __ 日"]
    for row, label, value in zip(table.rows, labels, values):
        shade(row.cells[0], GRAY)
        p1, p2 = row.cells[0].paragraphs[0], row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline_runs(p1, label, size=10)
        inline_runs(p2, value, size=10)
        for run in p1.runs:
            run.bold = True
        for cell in row.cells:
            set_cell_margins(cell, top=110, bottom=110)
    set_table_geometry(table, [2520, 6840])
    doc.add_page_break()


def add_image(doc, alt, relative_path):
    image_path = SOURCE.parent / relative_path
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image_path), width=Inches(6.18))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run(alt)
    set_font(run, size=9, italic=True, color=MUTED)


def parse_markdown(doc, source):
    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("> "):
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        image = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image:
            add_image(doc, image.group(1), image.group(2))
            i += 1
            continue
        if re.match(r"^#{2,4} ", line):
            level = len(line) - len(line.lstrip("#"))
            title = line[level + 1:]
            p = doc.add_paragraph(style={2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 3"))
            inline_runs(p, title, size={2: 16, 3: 13, 4: 11.5}.get(level, 11.5))
            i += 1
            continue
        if line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            table = doc.add_table(rows=1, cols=1)
            shade(table.cell(0, 0), "F5F7FA")
            p = table.cell(0, 0).paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("\n".join(code))
            set_font(run, name="Menlo", size=8.5, color=INK)
            set_table_geometry(table, [9360])
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[ :\-|]+\|$", lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[ :\-|]+\|$", lines[i]):
                    rows.append([part.strip() for part in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        bullet = re.match(r"^[-*] (.*)", line)
        numbered = re.match(r"^(\d+)\. (.*)", line)
        if bullet:
            add_bullet(doc, bullet.group(1))
            i += 1
            continue
        if numbered:
            add_bullet(doc, numbered.group(2), numbered=True)
            i += 1
            continue
        add_body_para(doc, line)
        i += 1


def main():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    parse_markdown(doc, SOURCE)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
